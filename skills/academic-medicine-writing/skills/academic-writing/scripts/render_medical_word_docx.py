#!/usr/bin/env python3
"""Render a medical manuscript Markdown source into a structured Word DOCX."""

from __future__ import annotations

import argparse
import html
import json
import re
import sys
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

PACKAGE_ROOT = Path(__file__).resolve().parents[3]
REVIEW_SCRIPT_DIR = PACKAGE_ROOT / "skills/academic-review/scripts"
if str(REVIEW_SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(REVIEW_SCRIPT_DIR))

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt
    DOCX_AVAILABLE = True
except ImportError:  # pragma: no cover - exercised in missing-dependency environments
    Document = None
    WD_ALIGN_PARAGRAPH = None
    Inches = None
    Pt = None
    DOCX_AVAILABLE = False

try:
    from audit_word_docx import audit_docx
except ImportError:  # pragma: no cover
    audit_docx = None


FIGURE_RE = re.compile(r"^\*\*Figure\s+(\d+)\.", re.IGNORECASE)
SUP_OR_BRACKET_CITE_RE = re.compile(r"(<sup>[^<]+</sup>|\[\d+(?:\s*(?:,|-)\s*\d+)*\])")
DOCX_AUDIT_PASS_TOKEN = "DOCX structure audit: PASS"
DEFAULT_REFERENCE_DOCX = PACKAGE_ROOT / "assets/templates/word/generic-medical-word-reference.docx"
DEFAULT_INLINE_PREVIEW_NAME = "review-preview-inline.docx"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
CONTENT_TYPES_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
RELS_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
PIC_NS = "http://schemas.openxmlformats.org/drawingml/2006/picture"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
FOOTER_REL_TYPE = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footer"
W = f"{{{W_NS}}}"
W_NS_MAP = {"w": W_NS}
DIRECT_STYLE_PROPS = {
    "Normal": {"sz": "24", "bold": False},
    "Title": {"sz": "24", "bold": True},
    "Heading1": {"sz": "24", "bold": True},
    "Heading2": {"sz": "24", "bold": True},
    "Caption": {"sz": "22", "bold": False},
}
PAGE_BREAK_BEFORE_HEADINGS = {"Key Points", "Abstract", "Introduction", "References", "Figure Legends", "Tables"}


def set_document_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    for name, size in (("Title", 12), ("Heading 1", 12), ("Heading 2", 12)):
        styles[name].font.name = "Times New Roman"
        styles[name].font.size = Pt(size)
        styles[name].font.bold = True
    styles["Caption"].font.name = "Times New Roman"
    styles["Caption"].font.size = Pt(11)


def split_markdown_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        line = lines[i].strip()
        cells = [cell.strip() for cell in line.strip("|").split("|")]
        if not all(set(cell) <= {"-", ":", " "} for cell in cells):
            rows.append(cells)
        i += 1
    return rows, i


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    width = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=width)
    table.style = "Table Grid"
    table.autofit = True
    for row_idx, row in enumerate(rows):
        for col_idx in range(width):
            cell = table.cell(row_idx, col_idx)
            cell.text = row[col_idx] if col_idx < len(row) else ""
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(8.5)
                    if row_idx == 0:
                        run.bold = True


def matching_figure(figures_dir: Path, number: str) -> Path | None:
    candidates = sorted(figures_dir.glob(f"figure{number}_*.png"))
    if not candidates:
        candidates = sorted(figures_dir.glob(f"fig{number}_*.png"))
    return candidates[0] if candidates else None


def add_paragraph_with_basic_bold(doc: Document, text: str, style: str | None = None):
    paragraph = doc.add_paragraph(style=style)
    if style == "Title":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run(part[2:-2] if part.startswith("**") and part.endswith("**") else part)
        if part.startswith("**") and part.endswith("**"):
            run.bold = True
        run.font.name = "Times New Roman"
        if style == "Caption":
            run.font.size = Pt(9)
    return paragraph


def render_with_python_docx(markdown: Path, output: Path, figures_dir: Path | None = None) -> int:
    lines = markdown.read_text(encoding="utf-8").splitlines()
    doc = Document()
    set_document_defaults(doc)
    expected_figures = 0
    i = 0
    while i < len(lines):
        raw = lines[i].rstrip()
        stripped = raw.strip()
        if not stripped:
            i += 1
            continue
        if stripped.startswith("|"):
            rows, i = split_markdown_table(lines, i)
            add_markdown_table(doc, rows)
            continue
        if stripped.startswith("# "):
            add_paragraph_with_basic_bold(doc, stripped[2:], "Title")
        elif stripped.startswith("## "):
            add_paragraph_with_basic_bold(doc, stripped[3:], "Heading 1")
        elif stripped.startswith("### "):
            add_paragraph_with_basic_bold(doc, stripped[4:], "Heading 2")
        else:
            fig_match = FIGURE_RE.match(stripped)
            style = "Caption" if fig_match else None
            add_paragraph_with_basic_bold(doc, stripped, style)
            if fig_match and figures_dir is not None:
                expected_figures += 1
                image = matching_figure(figures_dir, fig_match.group(1))
                if image is not None:
                    pic_para = doc.add_paragraph()
                    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = pic_para.add_run()
                    run.add_picture(str(image), width=Inches(6.0))
        i += 1

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    return expected_figures


def w_tag(local: str) -> str:
    return f"{W}{local}"


def ensure_child(parent: ET.Element, local: str) -> ET.Element:
    child = parent.find(f"w:{local}", W_NS_MAP)
    if child is None:
        child = ET.SubElement(parent, w_tag(local))
    return child


def find_docx_style(styles: ET.Element, style_id: str) -> ET.Element | None:
    for style in styles.findall(".//w:style", W_NS_MAP):
        if style.attrib.get(W + "styleId") == style_id:
            return style
    return None


def ensure_docx_style(styles: ET.Element, style_id: str, style_type: str, name: str) -> ET.Element:
    style = find_docx_style(styles, style_id)
    if style is None:
        style = ET.SubElement(styles, w_tag("style"))
        style.set(W + "styleId", style_id)
    style.set(W + "type", style_type)
    ensure_child(style, "name").set(W + "val", name)
    return style


def set_border_attrs(elem: ET.Element, sz: str = "8", color: str = "808080") -> None:
    elem.set(W + "val", "single")
    elem.set(W + "sz", sz)
    elem.set(W + "space", "0")
    elem.set(W + "color", color)


def ensure_border_set(parent: ET.Element, container: str, tags: tuple[str, ...], sz: str = "8") -> None:
    borders = ensure_child(parent, container)
    for tag in tags:
        set_border_attrs(ensure_child(borders, tag), sz=sz)


def set_style_font(style: ET.Element, size_half_points: str, bold: bool = False) -> None:
    rpr = ensure_child(style, "rPr")
    fonts = ensure_child(rpr, "rFonts")
    for attr in ("ascii", "hAnsi", "cs"):
        fonts.set(W + attr, "Times New Roman")
    ensure_child(rpr, "sz").set(W + "val", size_half_points)
    ensure_child(rpr, "szCs").set(W + "val", size_half_points)
    if bold:
        ensure_child(rpr, "b")
        ensure_child(rpr, "bCs")
    else:
        for tag in ("b", "bCs"):
            elem = rpr.find(f"w:{tag}", W_NS_MAP)
            if elem is not None:
                rpr.remove(elem)


def set_style_paragraph(
    style: ET.Element,
    before: str = "0",
    after: str = "120",
    align: str | None = None,
    keep_next: bool = False,
) -> None:
    ppr = ensure_child(style, "pPr")
    spacing = ensure_child(ppr, "spacing")
    spacing.set(W + "before", before)
    spacing.set(W + "after", after)
    spacing.set(W + "line", "276")
    spacing.set(W + "lineRule", "auto")
    if align:
        ensure_child(ppr, "jc").set(W + "val", align)
    if keep_next:
        ensure_child(ppr, "keepNext")


def set_table_grid_style(style: ET.Element) -> None:
    tbl_pr = ensure_child(style, "tblPr")
    ensure_border_set(tbl_pr, "tblBorders", ("top", "left", "bottom", "right", "insideH", "insideV"))
    cell_mar = ensure_child(tbl_pr, "tblCellMar")
    for side in ("top", "left", "bottom", "right"):
        margin = ensure_child(cell_mar, side)
        margin.set(W + "w", "80")
        margin.set(W + "type", "dxa")


def patch_styles_xml(styles_xml: bytes) -> bytes:
    ET.register_namespace("w", W_NS)
    styles = ET.fromstring(styles_xml)

    normal = ensure_docx_style(styles, "Normal", "paragraph", "Normal")
    set_style_font(normal, "24")
    set_style_paragraph(normal, after="120")

    title = ensure_docx_style(styles, "Title", "paragraph", "Title")
    set_style_font(title, "24", bold=True)
    set_style_paragraph(title, before="0", after="180", align="center")

    heading1 = ensure_docx_style(styles, "Heading1", "paragraph", "heading 1")
    set_style_font(heading1, "24", bold=True)
    set_style_paragraph(heading1, before="240", after="100", keep_next=True)

    heading2 = ensure_docx_style(styles, "Heading2", "paragraph", "heading 2")
    set_style_font(heading2, "24", bold=True)
    set_style_paragraph(heading2, before="160", after="80", keep_next=True)

    caption = ensure_docx_style(styles, "Caption", "paragraph", "caption")
    set_style_font(caption, "22")
    set_style_paragraph(caption, before="80", after="160")

    table_grid = ensure_docx_style(styles, "TableGrid", "table", "Table Grid")
    set_table_grid_style(table_grid)
    return ET.tostring(styles, encoding="utf-8", xml_declaration=True)


def run_properties_xml(style: str | None = None, bold: bool = False, size: str | None = None) -> str:
    props = DIRECT_STYLE_PROPS.get(style or "Normal", DIRECT_STYLE_PROPS["Normal"])
    effective_size = size or props["sz"]
    effective_bold = bold or bool(props["bold"])
    bold_xml = "<w:b/><w:bCs/>" if effective_bold else ""
    return (
        "<w:rPr>"
        '<w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:cs="Times New Roman"/>'
        f'{bold_xml}<w:sz w:val="{effective_size}"/><w:szCs w:val="{effective_size}"/>'
        "</w:rPr>"
    )


def paragraph_properties_xml(style: str | None = None, align: str | None = None) -> str:
    props = {
        "Title": {"before": "0", "after": "180", "align": "center", "keep": False},
        "Heading1": {"before": "240", "after": "100", "align": None, "keep": True},
        "Heading2": {"before": "160", "after": "80", "align": None, "keep": True},
        "Caption": {"before": "80", "after": "160", "align": None, "keep": False},
    }.get(style, {"before": "0", "after": "120", "align": None, "keep": False})
    jc = align or props["align"]
    parts = ["<w:pPr>"]
    if style:
        parts.append(f'<w:pStyle w:val="{style}"/>')
    parts.append(
        f'<w:spacing w:before="{props["before"]}" w:after="{props["after"]}" '
        'w:line="480" w:lineRule="auto"/>'
    )
    if jc:
        parts.append(f'<w:jc w:val="{jc}"/>')
    if props["keep"]:
        parts.append("<w:keepNext/>")
    parts.append("</w:pPr>")
    return "".join(parts)


def superscript_run_xml(text: str) -> str:
    return (
        f'<w:r><w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:cs="Times New Roman"/>'
        f'<w:vertAlign w:val="superscript"/><w:sz w:val="18"/><w:szCs w:val="18"/></w:rPr>'
        f'<w:t xml:space="preserve">{html.escape(text)}</w:t></w:r>'
    )


def text_runs_xml(text: str, style: str | None = None, size: str | None = None, base_bold: bool = False) -> str:
    chunks = []
    for token in SUP_OR_BRACKET_CITE_RE.split(text):
        if not token:
            continue
        if token.startswith("<sup>") and token.endswith("</sup>"):
            chunks.append(superscript_run_xml(token[5:-6]))
            continue
        if re.fullmatch(r"\[\d+(?:\s*(?:,|-)\s*\d+)*\]", token):
            chunks.append(superscript_run_xml(token[1:-1]))
            continue
        parts = re.split(r"(\*\*[^*]+\*\*)", token)
        for part in parts:
            if not part:
                continue
            bold = base_bold or (part.startswith("**") and part.endswith("**"))
            value = part[2:-2] if part.startswith("**") and part.endswith("**") else part
            chunks.append(
                f"<w:r>{run_properties_xml(style, bold=bold, size=size)}"
                f'<w:t xml:space="preserve">{html.escape(value)}</w:t></w:r>'
            )
    return "".join(chunks)


def paragraph_xml(
    text: str,
    style: str | None = None,
    align: str | None = None,
    page_break_before: bool = False,
) -> str:
    prefix = '<w:r><w:br w:type="page"/></w:r>' if page_break_before else ""
    return f"<w:p>{paragraph_properties_xml(style, align)}{prefix}{text_runs_xml(text, style)}</w:p>"


def legacy_text_runs_xml(text: str, style: str | None = None, size: str | None = None, base_bold: bool = False) -> str:
    chunks = []
    for part in re.split(r"(\*\*[^*]+\*\*)", text):
        if not part:
            continue
        bold = base_bold or (part.startswith("**") and part.endswith("**"))
        value = part[2:-2] if part.startswith("**") and part.endswith("**") else part
        chunks.append(
            f"<w:r>{run_properties_xml(style, bold=bold, size=size)}"
            f'<w:t xml:space="preserve">{html.escape(value)}</w:t></w:r>'
        )
    return "".join(chunks)


def table_borders_xml(container: str, tags: tuple[str, ...], sz: str = "6", color: str = "BFBFBF") -> str:
    parts = [f"<w:{container}>"]
    for tag in tags:
        parts.append(f'<w:{tag} w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>')
    parts.append(f"</w:{container}>")
    return "".join(parts)


def table_cell_paragraph_xml(text: str, header: bool = False) -> str:
    return (
        '<w:p><w:pPr><w:spacing w:before="0" w:after="0" w:line="240" w:lineRule="auto"/></w:pPr>'
        f'{text_runs_xml(text, None, size="18", base_bold=header)}</w:p>'
    )


def table_xml(rows: list[list[str]]) -> str:
    if not rows:
        return ""
    width = max(len(row) for row in rows)
    col_width = max(1000, min(2200, 9360 // max(width, 1)))
    grid_cols = "".join(f'<w:gridCol w:w="{col_width}"/>' for _ in range(width))
    out = [
        '<w:tbl><w:tblPr><w:tblStyle w:val="TableGrid"/><w:tblW w:w="9360" w:type="dxa"/>'
        '<w:tblLayout w:type="fixed"/>'
        + table_borders_xml("tblBorders", ("top", "left", "bottom", "right", "insideH", "insideV"))
        + '<w:tblCellMar><w:top w:w="80" w:type="dxa"/><w:left w:w="80" w:type="dxa"/>'
        '<w:bottom w:w="80" w:type="dxa"/><w:right w:w="80" w:type="dxa"/></w:tblCellMar>'
        "</w:tblPr>"
        f"<w:tblGrid>{grid_cols}</w:tblGrid>"
    ]
    for row_idx, row in enumerate(rows):
        out.append("<w:tr>")
        if row_idx == 0:
            out.append("<w:trPr><w:tblHeader/></w:trPr>")
        for col_idx in range(width):
            text = row[col_idx] if col_idx < len(row) else ""
            cell_pr = (
                f'<w:tcPr><w:tcW w:w="{col_width}" w:type="dxa"/>'
                + table_borders_xml("tcBorders", ("top", "left", "bottom", "right"))
                + ('<w:shd w:val="clear" w:color="auto" w:fill="F2F2F2"/>' if row_idx == 0 else "")
                + "</w:tcPr>"
            )
            out.append(f"<w:tc>{cell_pr}{table_cell_paragraph_xml(text, header=row_idx == 0)}</w:tc>")
        out.append("</w:tr>")
    out.append("</w:tbl>")
    return "".join(out)


def image_xml(rid: str, filename: str, docpr_id: int, width_emu: int = 5_486_400) -> str:
    height_emu = 3_657_600
    name = html.escape(filename)
    return f"""
<w:p><w:pPr><w:jc w:val="center"/></w:pPr><w:r><w:drawing>
<wp:inline distT="0" distB="0" distL="0" distR="0">
<wp:extent cx="{width_emu}" cy="{height_emu}"/>
<wp:effectExtent l="0" t="0" r="0" b="0"/>
<wp:docPr id="{docpr_id}" name="{name}"/>
<wp:cNvGraphicFramePr><a:graphicFrameLocks noChangeAspect="1"/></wp:cNvGraphicFramePr>
<a:graphic><a:graphicData uri="{PIC_NS}">
<pic:pic><pic:nvPicPr><pic:cNvPr id="0" name="{name}"/><pic:cNvPicPr/></pic:nvPicPr>
<pic:blipFill><a:blip r:embed="{rid}"/><a:stretch><a:fillRect/></a:stretch></pic:blipFill>
<pic:spPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="{width_emu}" cy="{height_emu}"/></a:xfrm>
<a:prstGeom prst="rect"><a:avLst/></a:prstGeom></pic:spPr></pic:pic>
</a:graphicData></a:graphic>
</wp:inline>
</w:drawing></w:r></w:p>"""


def footer_xml() -> bytes:
    return f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:ftr xmlns:w="{W_NS}">
  <w:p>
    <w:pPr><w:jc w:val="center"/></w:pPr>
    <w:r><w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:cs="Times New Roman"/><w:sz w:val="24"/><w:szCs w:val="24"/></w:rPr><w:fldChar w:fldCharType="begin"/></w:r>
    <w:r><w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:cs="Times New Roman"/><w:sz w:val="24"/><w:szCs w:val="24"/></w:rPr><w:instrText xml:space="preserve">PAGE</w:instrText></w:r>
    <w:r><w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:cs="Times New Roman"/><w:sz w:val="24"/><w:szCs w:val="24"/></w:rPr><w:fldChar w:fldCharType="separate"/></w:r>
    <w:r><w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:cs="Times New Roman"/><w:sz w:val="24"/><w:szCs w:val="24"/></w:rPr><w:t>1</w:t></w:r>
    <w:r><w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:cs="Times New Roman"/><w:sz w:val="24"/><w:szCs w:val="24"/></w:rPr><w:fldChar w:fldCharType="end"/></w:r>
  </w:p>
</w:ftr>'''.encode("utf-8")


def add_content_type_defaults(content_types: bytes, extensions: set[str], include_footer: bool = True) -> bytes:
    ET.register_namespace("", CONTENT_TYPES_NS)
    root = ET.fromstring(content_types)
    existing = {
        elem.attrib.get("Extension")
        for elem in root.findall(f"{{{CONTENT_TYPES_NS}}}Default")
    }
    for ext in sorted(extensions):
        if ext not in existing:
            elem = ET.SubElement(root, f"{{{CONTENT_TYPES_NS}}}Default")
            elem.set("Extension", ext)
            elem.set("ContentType", "image/png" if ext == "png" else "application/octet-stream")
    if include_footer:
        footer_part = "/word/footer1.xml"
        has_footer_override = any(
            elem.attrib.get("PartName") == footer_part
            for elem in root.findall(f"{{{CONTENT_TYPES_NS}}}Override")
        )
        if not has_footer_override:
            elem = ET.SubElement(root, f"{{{CONTENT_TYPES_NS}}}Override")
            elem.set("PartName", footer_part)
            elem.set("ContentType", "application/vnd.openxmlformats-officedocument.wordprocessingml.footer+xml")
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def document_rels_xml(image_rels: list[tuple[str, str]], include_footer: bool = True) -> bytes:
    ET.register_namespace("", RELS_NS)
    root = ET.Element(f"{{{RELS_NS}}}Relationships")
    for rid, target in image_rels:
        elem = ET.SubElement(root, f"{{{RELS_NS}}}Relationship")
        elem.set("Id", rid)
        elem.set("Type", "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image")
        elem.set("Target", target)
    if include_footer:
        elem = ET.SubElement(root, f"{{{RELS_NS}}}Relationship")
        elem.set("Id", "rIdFooter1")
        elem.set("Type", FOOTER_REL_TYPE)
        elem.set("Target", "footer1.xml")
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def render_with_reference_docx(
    markdown: Path,
    output: Path,
    reference_docx: Path,
    figures_dir: Path | None = None,
    figure_mode: str = "embed-preview",
) -> int:
    lines = markdown.read_text(encoding="utf-8").splitlines()
    body: list[str] = []
    media: list[tuple[str, Path, str]] = []
    expected_figures = 0
    docpr_id = 1
    i = 0
    while i < len(lines):
        stripped = lines[i].strip()
        if not stripped:
            i += 1
            continue
        if stripped.startswith("|"):
            rows, i = split_markdown_table(lines, i)
            body.append(table_xml(rows))
            continue
        if stripped.startswith("# "):
            body.append(paragraph_xml(stripped[2:], "Title", "center"))
        elif stripped.startswith("## "):
            heading = stripped[3:]
            body.append(paragraph_xml(heading, "Heading1", page_break_before=heading in PAGE_BREAK_BEFORE_HEADINGS))
        elif stripped.startswith("### "):
            body.append(paragraph_xml(stripped[4:], "Heading2"))
        else:
            fig_match = FIGURE_RE.match(stripped)
            body.append(paragraph_xml(stripped, "Caption" if fig_match else None))
            if fig_match and figures_dir is not None:
                expected_figures += 1
                image = matching_figure(figures_dir, fig_match.group(1))
                if image is not None and figure_mode == "embed-preview":
                    rid = f"rIdImage{expected_figures}"
                    name = f"figure{fig_match.group(1)}.png"
                    media.append((rid, image, name))
                    body.append(image_xml(rid, name, docpr_id))
                    docpr_id += 1
        i += 1

    document_xml = f'''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
  xmlns:r="{R_NS}" xmlns:wp="{WP_NS}" xmlns:a="{A_NS}" xmlns:pic="{PIC_NS}">
  <w:body>
    {''.join(body)}
    <w:sectPr><w:footerReference w:type="default" r:id="rIdFooter1"/><w:pgSz w:w="12240" w:h="15840"/><w:pgMar w:top="1440" w:right="1440" w:bottom="1440" w:left="1440"/></w:sectPr>
  </w:body>
</w:document>'''.encode("utf-8")

    output.parent.mkdir(parents=True, exist_ok=True)
    image_exts = {path.suffix.lower().lstrip(".") for _, path, _ in media}
    with zipfile.ZipFile(reference_docx) as src, zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as dst:
        for info in src.infolist():
            if info.filename in {"word/document.xml", "word/_rels/document.xml.rels"}:
                continue
            if info.filename == "[Content_Types].xml":
                dst.writestr(info, add_content_type_defaults(src.read(info.filename), image_exts))
            elif info.filename == "word/styles.xml":
                dst.writestr(info, patch_styles_xml(src.read(info.filename)))
            else:
                dst.writestr(info, src.read(info.filename))
        dst.writestr("word/document.xml", document_xml)
        dst.writestr("word/_rels/document.xml.rels", document_rels_xml([(rid, f"media/{name}") for rid, _, name in media]))
        dst.writestr("word/footer1.xml", footer_xml())
        for _, image_path, name in media:
            dst.writestr(f"word/media/{name}", image_path.read_bytes())
    return expected_figures


def render_markdown_to_docx(
    markdown: Path,
    output: Path,
    figures_dir: Path | None = None,
    reference_docx: Path | None = None,
    figure_mode: str = "embed-preview",
) -> int:
    if reference_docx is None and DEFAULT_REFERENCE_DOCX.exists():
        reference_docx = DEFAULT_REFERENCE_DOCX
    if DOCX_AVAILABLE and reference_docx is None:
        return render_with_python_docx(markdown, output, figures_dir)
    if reference_docx is not None:
        return render_with_reference_docx(markdown, output, reference_docx, figures_dir, figure_mode)
    raise SystemExit("python-docx is unavailable; provide --reference-docx or install python-docx")


def replace_prefixed_line(text: str, prefix: str, replacement: str) -> str:
    lines = text.splitlines()
    for index, line in enumerate(lines):
        if line.startswith(prefix):
            lines[index] = replacement
            return "\n".join(lines) + "\n"
    return text.rstrip() + "\n" + replacement + "\n"


def replace_any_prefixed_line(text: str, prefixes: tuple[str, ...], replacement: str) -> str:
    lines = text.splitlines()
    for index, line in enumerate(lines):
        if any(line.startswith(prefix) for prefix in prefixes):
            lines[index] = replacement
            return "\n".join(lines) + "\n"
    return text.rstrip() + "\n" + replacement + "\n"


def upsert_section(text: str, heading: str, body: str, before_heading: str | None = None) -> str:
    section = f"## {heading}\n\n{body.strip()}\n"
    pattern = re.compile(rf"\n## {re.escape(heading)}\n\n.*?(?=\n## |\Z)", re.DOTALL)
    if pattern.search(text):
        return pattern.sub("\n" + section, text).rstrip() + "\n"
    if before_heading and f"\n## {before_heading}\n" in text:
        return text.replace(f"\n## {before_heading}\n", "\n" + section + f"\n## {before_heading}\n", 1)
    return text.rstrip() + "\n\n" + section


def audit_brief(report: dict) -> str:
    metrics = report.get("metrics", {})
    status = report.get("status", "BLOCKED")
    status_text = DOCX_AUDIT_PASS_TOKEN if status == "PASS" else f"DOCX structure audit: {status}"
    return (
        f"{status_text} "
        f"(zip_entries={metrics.get('zip_entries', 0)}, styles={metrics.get('style_count', 0)}, "
        f"embedded_media={metrics.get('media_count', 0)}, drawings={metrics.get('drawings', 0)}, "
        f"table_style_coverage={metrics.get('table_style_coverage', 0)}, "
        f"cell_border_coverage={metrics.get('cell_border_coverage', 0)}, "
        f"table_grid_coverage={metrics.get('table_grid_coverage', 0)}, "
        f"double_spacing_coverage={metrics.get('double_spacing_coverage', 0)}, "
        f"has_page_numbers={str(metrics.get('has_page_numbers', False)).lower()}, "
        f"expected_figures={metrics.get('expected_figures', 0)}, "
        f"figure_mode={metrics.get('figure_mode', 'unknown')}, "
        f"page_breaks={metrics.get('page_break_count', 0)}, "
        f"superscript_refs={metrics.get('superscript_reference_runs', 0)}, "
        f"minimal_ooxml={str(metrics.get('minimal_ooxml', True)).lower()})"
    )


def expected_figure_count(metrics: dict) -> int:
    try:
        return int(metrics.get("expected_figures", 0) or 0)
    except (TypeError, ValueError):
        return 0


def visual_display_pass_line(metrics: dict, figure_mode: str) -> str:
    expected_figures = expected_figure_count(metrics)
    media_count = metrics.get("media_count", 0)
    drawings = metrics.get("drawings", 0)
    if expected_figures == 0:
        return (
            "- Visual display gate: NOT APPLICABLE. No figures are planned for this artifact; "
            "DOCX audit used expected_figures=0, so no `word/media` or `w:drawing` evidence is required."
        )
    if figure_mode == "upload-only":
        return (
            "- Visual display gate: PASS. separate figure upload route: "
            f"expected_figures={expected_figures}; the primary DOCX is submission-clean with "
            f"{media_count} `word/media` files and {drawings} `w:drawing` elements; inspect upload-ready "
            "assets under `paper/figures/` before submission."
        )
    return (
        "- Visual display gate: PASS. embedded preview route: planned figure callouts/legends are present, "
        f"expected_figures={expected_figures}, and the DOCX contains {media_count} `word/media` files and "
        f"{drawings} `w:drawing` elements."
    )


def visual_evidence_section_line(metrics: dict, figure_mode: str) -> str:
    expected_figures = expected_figure_count(metrics)
    media_count = metrics.get("media_count", 0)
    drawings = metrics.get("drawings", 0)
    if expected_figures == 0:
        return "- Embedded visual evidence: not applicable; expected_figures=0 and no figures are planned."
    if figure_mode == "upload-only":
        return (
            "- Embedded visual evidence: not required for the primary submission-clean DOCX; "
            f"figure_mode=upload-only, expected_figures={expected_figures}, embedded media={media_count}, "
            f"drawings={drawings}."
        )
    return (
        f"- Embedded visual evidence: {media_count} `word/media` files and "
        f"{drawings} `w:drawing` elements for expected_figures={expected_figures}."
    )


def audit_section_body(report: dict, audit_json: Path | None, package_dir: Path, figure_mode: str = "embed-preview") -> str:
    metrics = report.get("metrics", {})
    rel_audit = None
    if audit_json is not None:
        try:
            rel_audit = audit_json.resolve().relative_to(package_dir.resolve())
        except ValueError:
            rel_audit = audit_json
    reasons = report.get("reasons") or []
    return "\n".join(
        [
            f"- DOCX structure audit: {report.get('status', 'BLOCKED')}.",
            f"- Audit JSON: `{rel_audit or 'paper/docx-structure-audit.json'}`.",
            (
                f"- ZIP entries: {metrics.get('zip_entries', 0)}; style count: {metrics.get('style_count', 0)}; "
                f"paragraph style coverage: {metrics.get('paragraph_style_coverage', 0)}; "
                f"table style coverage: {metrics.get('table_style_coverage', 0)}."
            ),
            (
                f"- DOCX layout audit: cell border coverage {metrics.get('cell_border_coverage', 0)}; "
                f"table grid coverage {metrics.get('table_grid_coverage', 0)}; "
                f"double-space coverage {metrics.get('double_spacing_coverage', 0)}; "
                f"continuous page numbering {metrics.get('has_page_numbers', False)}; "
                f"page breaks {metrics.get('page_break_count', 0)}; "
                f"superscript reference runs {metrics.get('superscript_reference_runs', 0)}; "
                f"bracket numeric citations {metrics.get('bracket_numeric_citations', 0)}; "
                f"core style sizes {metrics.get('core_style_half_points', {})}."
            ),
            visual_evidence_section_line(metrics, figure_mode),
            (
                "- Required package parts: "
                f"`word/settings.xml`={metrics.get('has_word_settings_xml', False)}, "
                f"`word/styles.xml`={metrics.get('has_word_styles_xml', False)}, "
                f"`word/theme/`={metrics.get('has_word_theme', False)}, "
                f"`word/numbering.xml`={metrics.get('has_word_numbering_xml', False)}; "
                f"minimal OOXML={metrics.get('minimal_ooxml', True)}."
            ),
            f"- Blocker reasons: {'; '.join(reasons) if reasons else 'none'}.",
        ]
    )


def sync_submission_package(path: Path, report: dict, audit_json: Path | None, figure_mode: str = "embed-preview") -> None:
    if not path.exists():
        return
    text = path.read_text(encoding="utf-8")
    status = report.get("status", "BLOCKED")
    metrics = report.get("metrics", {})
    reasons = report.get("reasons") or []
    brief = audit_brief(report)
    if status == "PASS":
        format_line = (
            "- Format-specific production gate: PASS (word-first). `paper/manuscript.docx` and "
            f"`paper/manuscript.md` exist; no `main.pdf` is required; {brief}."
        )
        visual_line = visual_display_pass_line(metrics, figure_mode)
        table_line = (
            "- Table aesthetics gate: PASS for structural Word and DOCX layout audit. Editable Word tables "
            f"are present with table_style_coverage={metrics.get('table_style_coverage', 0)}, "
            f"cell_border_coverage={metrics.get('cell_border_coverage', 0)}, and "
            f"table_grid_coverage={metrics.get('table_grid_coverage', 0)}; final author/editorial-system "
            "visual inspection remains a pre-submission check."
        )
    else:
        reason_text = "; ".join(reasons) if reasons else "DOCX structure audit did not pass"
        format_line = f"- Format-specific production gate: BLOCKED (word-first). {brief}; reasons: {reason_text}."
        visual_line = f"- Visual display gate: BLOCKED for DOCX package. {brief}; reasons: {reason_text}."
        table_line = f"- Table aesthetics gate: BLOCKED. {brief}; reasons: {reason_text}."

    text = replace_prefixed_line(text, "- Format-specific production gate:", format_line)
    text = replace_prefixed_line(text, "- Visual display gate:", visual_line)
    text = replace_prefixed_line(text, "- Table aesthetics gate:", table_line)
    text = text.replace(
        "4. Word table visual inspection could not be performed in the current environment.",
        "4. Final author/editorial-system visual inspection of the uploaded DOCX and tables remains required.",
    )
    text = upsert_section(
        text,
        "DOCX Structure Audit",
        audit_section_body(report, audit_json, path.parent, figure_mode),
        before_heading="Main Artifacts",
    )
    path.write_text(text, encoding="utf-8")


def sync_review_report(path: Path, report: dict, audit_json: Path | None, figure_mode: str = "embed-preview") -> None:
    if not path.exists():
        return
    text = path.read_text(encoding="utf-8")
    metrics = report.get("metrics", {})
    status = report.get("status", "BLOCKED")
    if status == "PASS":
        verdict_line = (
            "Submission-ready: NO. Draft scaffold generated: YES, with blocking caveats. The package satisfies "
            "the word-first DOCX structure audit, but fails submission-readiness because core author statements, "
            "real clinical provenance, validated adjustment, and full reference development are missing."
        )
        issue_line = (
            "| production follow-up | DOCX structure audit passes; final editorial-system upload inspection is still required. "
            "| production check | manuscript.docx, paper/docx-structure-audit.json | Inspect the uploaded DOCX/figures in the journal system before submission. |"
        )
        visual_line = visual_display_pass_line(metrics, figure_mode)
    else:
        reason_text = "; ".join(report.get("reasons") or ["DOCX structure audit did not pass"])
        verdict_line = (
            "Submission-ready: NO. Draft scaffold generated: YES, with blocking caveats. The package fails "
            f"the word-first DOCX structure audit ({reason_text}) in addition to scientific and statement blockers."
        )
        issue_line = (
            f"| major compliance gap | DOCX structure audit is blocked: {reason_text}. | production weakness | "
            "manuscript.docx, paper/docx-structure-audit.json | Regenerate DOCX with normal Word package structure and embedded visual evidence. |"
        )
        visual_line = f"- Visual display gate: BLOCKED for DOCX package; {reason_text}."

    text = replace_prefixed_line(text, "Submission-ready:", verdict_line)
    text = replace_any_prefixed_line(
        text,
        (
            "| major compliance gap | Table aesthetics cannot be visually inspected in Word here.",
            "| production follow-up | DOCX structure audit passes;",
            "| major compliance gap | DOCX structure audit is blocked:",
        ),
        issue_line,
    )
    text = replace_prefixed_line(text, "- Visual display gate:", visual_line)
    text = upsert_section(
        text,
        "DOCX Structure Audit",
        audit_section_body(report, audit_json, path.parent, figure_mode),
        before_heading="Research Veto",
    )
    path.write_text(text, encoding="utf-8")


def sync_package_reports(package_dir: Path, report: dict, audit_json: Path | None = None, figure_mode: str = "embed-preview") -> None:
    sync_submission_package(package_dir / "submission-package.md", report, audit_json, figure_mode)
    sync_review_report(package_dir / "review-report.md", report, audit_json, figure_mode)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("markdown", type=Path)
    parser.add_argument("output", type=Path)
    parser.add_argument("--figures-dir", type=Path)
    parser.add_argument(
        "--reference-docx",
        dest="reference_docx",
        type=Path,
        help=f"Reference DOCX to preserve Word package structure; defaults to {DEFAULT_REFERENCE_DOCX}",
    )
    parser.add_argument("--audit-json", type=Path)
    parser.add_argument("--sync-package-reports", action="store_true")
    parser.add_argument(
        "--figure-mode",
        choices=("upload-only", "embed-preview"),
        default="embed-preview",
        help=f"Use upload-only for submission-clean DOCX or embed-preview for reader preview DOCX such as {DEFAULT_INLINE_PREVIEW_NAME}.",
    )
    args = parser.parse_args()

    expected = render_markdown_to_docx(
        args.markdown,
        args.output,
        args.figures_dir,
        args.reference_docx,
        args.figure_mode,
    )
    if args.audit_json or args.sync_package_reports:
        if audit_docx is None:
            raise SystemExit("cannot import audit_word_docx")
        report = audit_docx(
            args.output,
            expected_figures=expected,
            allow_separate_figures=args.figure_mode == "upload-only",
            forbid_embedded_figures=args.figure_mode == "upload-only",
        )
        report.setdefault("metrics", {})["expected_figures"] = expected
        report.setdefault("metrics", {})["figure_mode"] = args.figure_mode
        if args.audit_json:
            args.audit_json.write_text(json.dumps(report, indent=2, sort_keys=True) + "\n", encoding="utf-8")
        if args.sync_package_reports:
            sync_package_reports(args.output.parent, report, args.audit_json, args.figure_mode)
        if report["status"] != "PASS":
            return 1
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
